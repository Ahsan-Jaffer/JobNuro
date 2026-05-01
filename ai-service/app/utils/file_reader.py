import os
import fitz
import docx
from fastapi import UploadFile, HTTPException

ALLOWED_CONTENT_TYPES = {
    "application/pdf": "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}


def validate_resume_file(file: UploadFile, max_size_mb: int) -> str:
    if not file:
        raise HTTPException(status_code=400, detail="Resume file is required")

    file_type = ALLOWED_CONTENT_TYPES.get(file.content_type)

    if not file_type:
        raise HTTPException(
            status_code=400,
            detail="Only PDF and DOCX resume files are allowed",
        )

    return file_type


async def save_upload_file(file: UploadFile, destination_path: str, max_size_mb: int) -> None:
    max_size_bytes = max_size_mb * 1024 * 1024
    total_size = 0

    with open(destination_path, "wb") as buffer:
        while True:
            chunk = await file.read(1024 * 1024)
            if not chunk:
                break

            total_size += len(chunk)

            if total_size > max_size_bytes:
                buffer.close()
                if os.path.exists(destination_path):
                    os.remove(destination_path)

                raise HTTPException(
                    status_code=413,
                    detail=f"Resume file size cannot exceed {max_size_mb} MB",
                )

            buffer.write(chunk)


def extract_text_from_pdf(file_path: str) -> str:
    text_parts = []

    with fitz.open(file_path) as pdf_document:
        for page in pdf_document:
            text_parts.append(page.get_text())

    return "\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    document = docx.Document(file_path)

    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]

    table_text = []
    for table in document.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                table_text.append(" | ".join(row_text))

    return "\n".join(paragraphs + table_text)


def extract_text_from_resume(file_path: str, file_type: str) -> str:
    if file_type == "pdf":
        return extract_text_from_pdf(file_path)

    if file_type == "docx":
        return extract_text_from_docx(file_path)

    raise HTTPException(status_code=400, detail="Unsupported resume file type")